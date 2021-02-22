import os
import datetime
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
import pandas as pd
import numpy as np

import cfg
from libs.RC.getInfo_functions import get_INFO_posRes, get_INFO_decon

genotypes = ['CARRIER', 'CARRIER-Georgian', 'CARRIER-Georgian-Problem',
            'CARRIER - With soft-clipped reads', 'CARRIER-Problem - With soft-clipped reads',
            'CARRIER-Druze', 'CARRIER-Druze-Problem', 'CARRIER-Problem', 'HOM'] # For createReports().
cnvs = ['CNV', 'CNV Big Del Boundaries Different as Reported', 'CNV-Problem', 'CNV-Problem Big Del Boundaries Different as Reported'] # For createReports().
cnvs_problem = ['CNV-Problem', 'CNV-Problem Big Del Boundaries Different as Reported'] # For createReports().

REPORTS_folder = 'REPORTS' # Folder for output.
LC = 'Low Confidence' # Text for problem.
GC = 'Pathogenic in Georgian ethnicity' # Text for NCF1 Georgian.
DC = 'Pathogenic in Druze' # Text for HOGA1 Druze.
agid_gc = 'AG1470' # Georgian AGID
agid_dc = 'AG2482' # Druze AGID
NC = '_NC'

HETOFFICEOLD = '*נשאות באלל אחד )הטרוזיגוט('
HETOFFICENEW = 'נשאות באלל אחד *(הטרוזיגוט)'
HOMOFFICEOLD = '*נשאות בשני אללים )הומוזיגוט('
HOMOFFICENEW = 'נשאות בשני אללים *(הומוזיגוט)'

GQXTOP = 15 # ['GQX'] < GQXTOP = PROBLEM
AFVTOP = 30 # ['Alt Variant Freq'] > AFVBOTTOM and ['Alt Variant Freq'] < AFVTOP = PROBLEM
AFVBOTTOM = 10



"""
    Creates initial table.
    Used in:
        - createReports()
    INPUT:
        1) hdr_cells - the template for the docx table.
    *NO OUTPUT - edits in place*
"""
def create_Table(hdr_cells):
    hdr_cells[0].paragraphs[0].add_run('סטטוס')
    hdr_cells[1].paragraphs[0].add_run('מוטציה')
    hdr_cells[2].paragraphs[0].add_run('גן')
    hdr_cells[3].paragraphs[0].add_run('מחלה')
    for i in range(4): # Change header formatting.
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.name = 'Arial'
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(12)


"""
    Updates table.
    Used in:
        - createReports()
    INPUT:
        1) row_cells - template for current row in table.
        2) status - mutation status info.
        3) mutation - mutation name.
        4) gene - mutation gene.
        5) CA - mutation disease.
        6) problem - Boolean, False = regular mutation / True = problematic mutation, needs to be red.
    *NO OUTPUT - edits in place*
"""
def update_Table(row_cells, status, mutation, gene, CA, problem=False):
    row_cells[0].paragraphs[0].add_run(status)
    row_cells[1].text = mutation
    row_cells[2].text = gene
    row_cells[3].text = CA
    for i in range(4): # Change formatting.
        row_cells[i].paragraphs[0].runs[0].font.name = 'Arial'
        row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    if problem:
        row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)


"""
    Creates the actual reports (docx documents) based on all the data gathered and the templates provided.
    Used in:
        - createReports()
    INPUT:
        1) sample - current sample.
        2) sampleStatus - dictionary with KEY = sample / VALUE = [sample's status, geno indices, cnv indices].
        3) posResFiltered - filtered pd dataframe of genotyping results data.
        4) decon_filtered - filtered pd dataframe of CNV results data.
        5) sampleInfoTable - pd dataframe of sample extra information data.
        6) office - Boolean, True = office version less than 2016 / False = office version 2016 or higher.
        7) output_folder - full path to the results folder.
        8) MyScreen_version - the program version and version date.
    *NO OUTPUT - creates the reports*
"""
def createReports(sample, sampleStatus, posResFiltered, decon_filtered,
        sampleInfoTable, office, output_folder, MyScreen_version, logger_name,
        hospital, paths):

    template_sick = cfg.template_sick.format(hospital)
    template_norm = cfg.template_norm.format(hospital)
    template_carrier = cfg.template_carrier.format(hospital)
    template_carrier_and_sick = cfg.template_carrier_and_sick.format(hospital)

    D = {name : code for name, code in cfg.Panels}
    p = paths[sample][1]
    c = D.get(p)
    MyScreen_version = "{} \nTest Name:{} Test Code:{}".format(MyScreen_version, p, c)

    ### ------------- Pre-processing ------------- ###
    now = datetime.datetime.now()
    date = now.strftime("%d-%m-%Y") # Today's date.
    current_status = sampleStatus[sample][0] # Get status.
    geno_indices = sampleStatus[sample][1] # Get indices of occurences in 'positiveResults.tsv'.
    cnv_indices = sampleStatus[sample][2] # Get indices of occurences in 'DECoN_results.tsv'.

    ### ------------- Table creation ------------- ###
    classes = []
    for i in geno_indices:
        CA, gene, agid, mutation, mutation_gDNA, Classification, Genotype, GQX, AVF, RD, ARD, AD, moh, eth, DM, MM = get_INFO_posRes(i, posResFiltered, logger_name)
        classes.append(Classification)
    s = list(set(classes))
    if len(s) == 1 and s[0] == "NO_CALL-Problem":
        geno_indices = []
    # WT #
    genes = []
    for i in cnv_indices:
        CA, gene, agid, mutation, Classification, Genotype, Correlation, Ncomp, first, last, BF, expected, observed, ratio, moh, eth, DM, MM = get_INFO_decon(i, decon_filtered, logger_name)
        genes.append(gene)

    if (geno_indices==[]) and (cnv_indices==[]): # Sample does not exist in 'positiveResults.tsv' and has no DECoN mutations:
        document = Document(template_norm) # Template document.
        pVer = document.paragraphs[15].add_run(MyScreen_version) # Update version.
        pVer.font.name = 'Arial'
        pVer.font.size = Pt(10)

    elif current_status == "NORM": # Only NO-CALLS or WT or  or CONVector_NC.
        document = Document(template_norm) # Template document.
        pVer = document.paragraphs[15].add_run(MyScreen_version) # Update version.
        pVer.font.name = 'Arial'
        pVer.font.size = Pt(10)

    elif geno_indices==[] and all(g == "DMD" for g in genes):
        document = Document(template_norm) # Template document.
        pVer = document.paragraphs[15].add_run(MyScreen_version) # Update version.
        pVer.font.name = 'Arial'
        pVer.font.size = Pt(10)

    # CARRIER #
    elif current_status == "CARRIER": # Only HET.
        document = Document(template_carrier) # Template document.
        pVer = document.paragraphs[17].add_run(MyScreen_version) # Update version.
        pVer.font.name = 'Arial'
        pVer.font.size = Pt(10)
        # Create table:
        table = document.add_table(1, 4, style='Table Grid')
        hdr_cells = table.rows[0].cells
        create_Table(hdr_cells)
        # Insert mutations to table:
        for i in geno_indices:
            CA, gene, agid, mutation, mutation_gDNA, Classification, Genotype, GQX, AVF, RD, ARD, AD, moh, eth, DM, MM = get_INFO_posRes(i, posResFiltered, logger_name)
            if (int(posResFiltered.iloc[i]['GQX']) < GQXTOP) or ((int(posResFiltered.iloc[i]['Alt Variant Freq']) > AFVBOTTOM) and (int(posResFiltered.iloc[i]['Alt Variant Freq']) < AFVTOP)):
                problem = LC
                red = True
                if agid == agid_gc:
                    problem = "{} - {}".format(LC, GC)
                if agid == agid_dc:
                    problem = "{} - {}".format(LC, DC)
            elif agid == agid_gc:
                problem = GC
                red = True
            elif agid == agid_dc:
                problem = DC
                red = True
            else:
                problem = ""
                red = False
            if Classification not in genotypes:
                continue # skip table update.
            if office:
                status = HETOFFICEOLD
            else:
                status = HETOFFICENEW
            row_cells = table.add_row().cells
            update_Table(row_cells, status + '\n' + problem, mutation_gDNA, gene, CA, red)
        for i in cnv_indices:
            CA, gene, agid, mutation, Classification, Genotype, Correlation, Ncomp, first, last, BF, expected, observed, ratio, moh, eth, DM, MM = get_INFO_decon(i, decon_filtered, logger_name)
            amp = decon_filtered.iloc[i]['Annotation1'] # Get amplicon name.
            class_decon = decon_filtered.iloc[i]['Classification']
            if class_decon not in cnvs or NC in amp:
                continue
            if NC in amp:
                continue # skip row update.
            if gene == "DMD":
                continue
            if office:
                status = HETOFFICEOLD
            else:
                status = HETOFFICENEW
            if class_decon in cnvs_problem:
                problem = LC
                red = True
            else:
                problem = ""
                red = False
            row_cells = table.add_row().cells
            update_Table(row_cells, status + '\n' + problem, mutation, gene, CA)
        # Move table to middle of document:
        para = document.paragraphs[13] # Location.
        p_obj = para._p
        t_obj = table._tbl
        p_obj.addnext(t_obj)

    # SICK #
    elif current_status == "SICK": # Only HOM.
        document = Document(template_sick) # Template document.
        pVer = document.paragraphs[19].add_run(MyScreen_version) # Update version.
        pVer.font.name = 'Arial'
        pVer.font.size = Pt(10)
        # Create table:
        table = document.add_table(1, 4, style='Table Grid')
        hdr_cells = table.rows[0].cells
        create_Table(hdr_cells)
        # Insert mutations to table:
        for i in geno_indices:
            CA, gene, agid, mutation, mutation_gDNA, Classification, Genotype, GQX, AVF, RD, ARD, AD, moh, eth, DM, MM = get_INFO_posRes(i, posResFiltered, logger_name)
            if (int(posResFiltered.iloc[i]['GQX']) < GQXTOP) or ((int(posResFiltered.iloc[i]['Alt Variant Freq']) > AFVBOTTOM) and (int(posResFiltered.iloc[i]['Alt Variant Freq']) < AFVTOP)):
                problem = LC
                red = True
                if agid == agid_gc:
                    problem = "{} - {}".format(LC, GC)
                if agid == agid_dc:
                    problem = "{} - {}".format(LC, DC)
            elif agid == agid_gc:
                problem = GC
                red = True
            elif agid == agid_dc:
                problem = DC
                red = True
            else:
                problem = ""
                red = False
            if Classification not in genotypes:
                continue # skip table update.
            if office:
                status = HOMOFFICEOLD
            else:
                status = HOMOFFICENEW
            row_cells = table.add_row().cells
            update_Table(row_cells, status + '\n' + problem, mutation_gDNA, gene, CA)
        for i in cnv_indices:
            CA, gene, agid, mutation, Classification, Genotype, Correlation, Ncomp, first, last, BF, expected, observed, ratio, moh, eth, DM, MM = get_INFO_decon(i, decon_filtered, logger_name)
            if class_decon not in cnvs or NC in amp:
                continue # skip table update.
            if gene == "DMD":
                continue
            if office:
                status = HOMOFFICEOLD
            else:
                status = HOMOFFICENEW
            if class_decon in cnvs_problem:
                problem = LC
                red = True
            else:
                problem = ""
                red = False
            row_cells = table.add_row().cells
            update_Table(row_cells, status + '\n' + problem, mutation, gene, CA)
        # Move table to middle of document:
        para = document.paragraphs[13] # Location.
        p_obj = para._p
        t_obj = table._tbl
        p_obj.addnext(t_obj)

    # CARRIER & SICK #
    else: # Is both HET and HOM.
        document = Document(template_carrier_and_sick) # Template document.
        pVer = document.paragraphs[20].add_run(MyScreen_version) # Update version.
        pVer.font.name = 'Arial'
        pVer.font.size = Pt(10)
        # Create hom table:
        hom_table = document.add_table(1, 4, style='Table Grid')
        hom_cells = hom_table.rows[0].cells
        create_Table(hom_cells)
        # Create het table:
        het_table = document.add_table(1, 4, style='Table Grid')
        het_cells = het_table.rows[0].cells
        create_Table(het_cells)
        # Insert mutations to table:
        for i in geno_indices:
            CA, gene, agid, mutation, mutation_gDNA, Classification, Genotype, GQX, AVF, RD, ARD, AD, moh, eth, DM, MM = get_INFO_posRes(i, posResFiltered, logger_name)
            if (int(posResFiltered.iloc[i]['GQX']) < GQXTOP) or ((int(posResFiltered.iloc[i]['Alt Variant Freq']) > AFVBOTTOM) and (int(posResFiltered.iloc[i]['Alt Variant Freq']) < AFVTOP)):
                problem = LC
                red = True
                if agid == agid_gc:
                    problem = "{} - {}".format(LC, GC)
                if agid == agid_dc:
                    problem = "{} - {}".format(LC, DC)
            elif agid == agid_gc:
                problem = GC
                red = True
            elif agid == agid_dc:
                problem = DC
                red = True
            else:
                problem = ""
                red = False
            if Classification not in genotypes:
                continue # skip table update.
            elif Genotype == 'het': # het mutation.
                if office:
                    status = HETOFFICEOLD
                else:
                    status = HETOFFICENEW
                row_cells = het_table.add_row().cells
                update_Table(row_cells, status + '\n' + problem, mutation_gDNA, gene, CA)
            else: # hom mutation.
                if office:
                    status = HOMOFFICEOLD
                else:
                    status = HOMOFFICENEW
                row_cells = hom_table.add_row().cells
                update_Table(row_cells, status + '\n' + problem, mutation_gDNA, gene, CA)
        for i in cnv_indices:
            CA, gene, agid, mutation, Classification, Genotype, Correlation, Ncomp, first, last, BF, expected, observed, ratio, moh, eth, DM, MM = get_INFO_decon(i, decon_filtered, logger_name)
            amp = decon_filtered.iloc[i]['Annotation1'] # Get amplicon name.
            class_decon = decon_filtered.iloc[i]['Classification']
            if class_decon not in cnvs or NC in amp:
                continue # skip table update.
            if gene == "DMD":
                continue
            if office:
                status = HOMOFFICEOLD
            else:
                status = HOMOFFICENEW
            if class_decon in cnvs_problem:
                problem = LC
                red = True
            else:
                problem = ""
                red = False
            row_cells = het_table.add_row().cells
            update_Table(row_cells, status + '\n' + problem, mutation, gene, CA)
        # Move hom table to middle of document:
        para1 = document.paragraphs[13] # Location.
        p_obj1 = para1._p
        t_obj1 = hom_table._tbl
        p_obj1.addnext(t_obj1)
        # Move het table to middle of document:
        para2 = document.paragraphs[16] # Location.
        p_obj2 = para2._p
        t_obj2 = het_table._tbl
        p_obj2.addnext(t_obj2)

    ### ------------- Personal info ------------- ###
    # Add current date, sample ID, ID, patient name, sample source and more:
    p0 = document.paragraphs[0].add_run(date)
    p0.font.name = 'Arial'
    p0.font.size = Pt(12)
    p1 = document.paragraphs[1].add_run(sample)
    p1.font.name = 'Arial'
    p1.font.size = Pt(12)
    ID = sampleInfoTable.at[sample, 'ID']
    if type(ID) != float and type(ID) != np.float64:
        p2 = document.paragraphs[2].add_run(ID)
        p2.font.name = 'Arial'
        p2.font.size = Pt(12)
    name = sampleInfoTable.at[sample, 'name']
    if type(name) != float and type(name) != np.float64:
        p3 = document.paragraphs[3].add_run(name)
        p3.font.name = 'Arial'
        p3.font.size = Pt(12)
    gender = sampleInfoTable.at[sample, 'gender']
    if type(gender) != float and type(gender) != np.float64:
        p4 = document.paragraphs[4].add_run(gender)
        p4.font.name = 'Arial'
        p4.font.size = Pt(12)
    street = sampleInfoTable.at[sample, 'street']
    if type(street) != float and type(street) != np.float64:
        p5 = document.paragraphs[5].add_run(street)
        p5.font.name = 'Arial'
        p5.font.size = Pt(12)
        space = 'א' # Take on Hebrew letter.
        pS = document.paragraphs[5].add_run(space) # Put in between street and city.
        pS.font.name = 'Arial'
        pS.font.size = Pt(12)
        pS.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # Color letter white to simulate blank space.
    city = sampleInfoTable.at[sample, 'city']
    if type(city) != float and type(city) != np.float64:
        p6 = document.paragraphs[5].add_run(city)
        p6.font.name = 'Arial'
        p6.font.size = Pt(12)
    phone = sampleInfoTable.at[sample, 'phone']
    if type(phone) != float and type(phone) != np.float64:
        p7 = document.paragraphs[6].add_run(phone)
        p7.font.name = 'Arial'
        p7.font.size = Pt(12)
    eth_m = sampleInfoTable.at[sample, 'eth_m']
    if type(eth_m) != float and type(eth_m) != np.float64:
        p8 = document.paragraphs[7].add_run(eth_m)
        p8.font.name = 'Arial'
        p8.font.size = Pt(12)
    eth_f = sampleInfoTable.at[sample, 'eth_f']
    if type(eth_f) != float and type(eth_f) != np.float64:
        p9 = document.paragraphs[8].add_run(eth_f)
        p9.font.name = 'Arial'
        p9.font.size = Pt(12)
    source = sampleInfoTable.at[sample, 'source']
    if type(source) != float and type(source) != np.float64:
        p10 = document.paragraphs[9].add_run(source)
        p10.font.name = 'Arial'
        p10.font.size = Pt(12)
    partner = sampleInfoTable.at[sample, 'partner']
    if type(partner) != float and type(partner) != np.float64:
        p11 = document.paragraphs[10].add_run(partner)
        p11.font.name = 'Arial'
        p11.font.size = Pt(12)
    if type(ID) == float or type(ID) == np.float64:
        ID = ""
    serial = sample + '_' + ID + '_' + date
    pSer = document.add_paragraph(serial)
    pSer.runs[0].font.name = 'Arial'
    pSer.runs[0].font.size = Pt(8)

    ### ------------- Save document ------------- ###
    # Save document with relevant name:
    output_doc = '\\XX_YY.docx' # Output document name format.
    doc_name = output_doc.replace('XX', sample) # Insert sample.
    doc_name = doc_name.replace('YY', ID) # Insert ID.
    folder_name = REPORTS_folder
    output_path = os.path.join(output_folder, folder_name)
    if not os.path.exists(output_path): # Create folder (if doesn't exist).
        os.makedirs(output_path)
    document.save(output_path + doc_name)
